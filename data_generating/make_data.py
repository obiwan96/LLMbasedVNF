from openai import OpenAI
from secret import OPENAI_API_KEY, KUBER_PASSWORD
from prompts import *
import re
import os
import argparse


openai_client = OpenAI(api_key=OPENAI_API_KEY)
data_path='data/'

def make_ansible(success_count = 1):
    # Don't use.
    for prompt in prompts['ansible']:
        response = openai_client.chat.completions.create(
            model = 'gpt-4o-mini', messages=[
                {"role" : "system", "content" : 'You are an expert in Kubernetes management.'},
                {"role" : "user", "content" : prompt}
                ], temperature=0)
        content=response.choices[0].message.content
        #print(content)
        #print('-------------------')
        yaml_pattern = r"```yaml(.*?)```"
        yaml_match = re.search(yaml_pattern, content, re.DOTALL)
        if yaml_match:
            yaml_content = yaml_match.group(1).strip()
            print("Extracted YAML content:")
            #print(yaml_content)

            yaml_file_path = data_path+"kubernetes_setup_"+str(success_count)+".yaml"
            with open(yaml_file_path, "w") as yaml_file:
                yaml_file.write(yaml_content)
            
            print(f"YAML content has been saved to {yaml_file_path}")
            
            # Ansible Playbook 실행
            import subprocess
            
            ansible_command = ["ansible-playbook", yaml_file_path, '--extra-vars', '"ansible_sudo_pass='+KUBER_PASSWORD+'"']
            
            try:
                result = subprocess.run(ansible_command, capture_output=True, text=True)
                if result.returncode == 0:
                    print("Ansible Playbook executed successfully.")
                    success_count += 1
                else:
                    print("Ansible Playbook execution failed.")
                    print("Error details:", result.stderr)
                    
                # 출력 결과
                print("Output:", result.stdout)
            except Exception as e:
                print("Failed to execute Ansible Playbook.")
                print(str(e))
        else:
            print("YAML content not found.")

    print(f'Total {success_count-1} YAML files have been created and executed successfully.')

def make_mop(args):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    example_mop_path='data/Example/'
    example_mop_list = os.listdir(example_mop_path)
    example_mop='Here are the example MOP files.\n'
    for example_mop_file_name in example_mop_list:
        #if (args.OpenStack and 'OpenStack' in example_mop_file_name) or (args.K8S and 'K8S' in example_mop_file_name):
        doc = Document(example_mop_path+example_mop_file_name)
        for para in doc.paragraphs:
            example_mop+=para.text+'\n'
    total_num_file=0
    langs = []
    if args.en:
        langs.append('en')
    if args.ko:
        langs.append('ko')
    if args.OpenStack:
        system_container= ['OpenStack', 'VM']
    elif args.K8S:
        system_container= ['Kubernetes', 'container']
    else:
        #Make intergration MOPs.
        system_container=['OpenStack/Kubernetes', 'VM/container']
    for lang in ['en', 'ko']:
        system, container = system_container
        for function in function_list:
            last_num=1
            for additional_command in additional_command_list[function]:
                for prompt in prompts['mop']:
                    formatted_prompt = example_mop+prompt.format(system=system, container=container, 
                        function=function, additional_command=additional_command)
                    if function=='firewall':
                        formatted_prompt+=' Use iptables operation, not ufw.'
                    if lang=='ko':
                        formatted_prompt+='Please write in Korean'
                    else:
                        formatted_prompt+='Please write in English'
                    if system=='OpenStack':
                        formatted_prompt+='Also, do not use the GUI(Horizon), use the CLI. '+ \
                            'Instead of setting floating IP on the created VM, use the Jump Host, '+ \
                            'which can connect to the internal VM, to connect to the newly created VM with SSH '+ \
                            'and operate the shell commands in SSH connection. To do this, enable SSH access through the password. \n' + \
                            "Don't make security groups or keypairs."
                    elif system=='OpenStack/Kubernetes':
                        formatted_prompt+= "Don't use GUI in the MOPs. Also, when using OpenStack, indicate user to don't make floating IP and using jump host."
                    response = openai_client.chat.completions.create(
                        model = 'gpt-4o-mini', messages=[
                            {"role" : "system", "content" : f'You are an expert in {system} management.'},
                            {"role" : "user", "content" : formatted_prompt}
                            ], temperature=0)
                    content=response.choices[0].message.content
                    #print(content)
                    #print('-------------------')
                    mop_file_path = data_path+f"{function}_setup_{additional_command[0]}_{lang}_{last_num}.docx"
                    last_num+=1
                    doc = Document()
                    for line in content.split('\n'):
                        if line:
                            if line.startswith('#'):
                                head_num=line.count('#')
                                doc.add_heading(line[head_num:], level=head_num)
                            else:
                                doc.add_paragraph(line)
                    doc.save(mop_file_path)
                    total_num_file+=1

    print(f'Total {total_num_file} doc files have been created.')

if __name__ == '__main__':
    argparser = argparse.ArgumentParser()
    argparser.add_argument('--OpenStack', action='store_true')
    argparser.add_argument('--K8S', action='store_true') 
    # Don't use K8S and OpenStack at the same time.
    # I'll asume that the user will use only one of them.
    argparser.add_argument('--en', action='store_true')
    argparser.add_argument('--ko', action='store_true')
    make_mop(argparser.parse_args())